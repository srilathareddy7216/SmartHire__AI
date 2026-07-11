"""
Extracts raw text from an uploaded resume file (PDF, DOCX, or TXT) so it can
be fed into the same cleaning / classification / recommendation pipeline
used for the training resumes.
"""
from __future__ import annotations

import io


def parse_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    # also pull text out of tables (many resumes use table layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def parse_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch on file extension and return extracted plain text."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext == "docx":
        return parse_docx(file_bytes)
    elif ext in ("txt", "text"):
        return parse_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Please upload PDF, DOCX, or TXT.")
